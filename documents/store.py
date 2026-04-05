"""
Document Store — upload and manage user documents in PostgreSQL.

Supported document types:
  - CV (PDF or DOCX)
  - Cover letter templates
  - University certificates / degrees
  - Work references / employer letters
  - Other supporting documents

All documents have their text extracted and embedded for RAG queries.
"""
from __future__ import annotations

import base64
import mimetypes
from pathlib import Path
from typing import List, Optional
from uuid import UUID

from loguru import logger
from sqlalchemy import select

from core.database import get_session
from core.models import Document, DocumentType
from matching.embedder import embed_text


class DocumentStore:
    """Manages all user documents — upload, retrieve, search."""

    async def upload(
        self,
        path: str | Path,
        doc_type: DocumentType,
        name: Optional[str] = None,
    ) -> Document:
        """
        Upload a document to the database.
        Extracts text content and generates embedding.
        """
        path = Path(path)
        if not path.exists():
            raise FileNotFoundError(f"Document not found: {path}")

        doc_name = name or path.stem
        logger.info(f"[DocStore] Uploading '{doc_name}' ({doc_type.value})")

        # Extract text based on file type
        content_text = self._extract_text(path)

        # Generate embedding
        embedding = None
        if content_text:
            try:
                embedding = embed_text(content_text[:2000])
            except Exception as e:
                logger.warning(f"[DocStore] Embedding failed: {e}")

        # Encode binary content
        content_b64 = base64.b64encode(path.read_bytes()).decode("utf-8")

        async with get_session() as session:
            doc = Document(
                doc_type=doc_type,
                name=doc_name,
                filename=path.name,
                content_text=content_text,
                content_bytes=content_b64,
                embedding=embedding,
                metadata_={
                    "file_size": path.stat().st_size,
                    "mime_type": mimetypes.guess_type(str(path))[0],
                    "original_path": str(path),
                },
            )
            session.add(doc)
            await session.flush()
            await session.refresh(doc)

        logger.info(f"[DocStore] Stored document id={doc.id}")
        return doc

    async def get_all(self, doc_type: Optional[DocumentType] = None) -> List[Document]:
        """Retrieve all documents, optionally filtered by type."""
        async with get_session() as session:
            query = select(Document)
            if doc_type:
                query = query.where(Document.doc_type == doc_type)
            result = await session.execute(query.order_by(Document.uploaded_at.desc()))
            return list(result.scalars().all())

    async def get_by_id(self, doc_id: UUID) -> Optional[Document]:
        async with get_session() as session:
            return await session.get(Document, doc_id)

    async def get_cv_text(self) -> Optional[str]:
        """Return the text content of the most recently uploaded CV."""
        docs = await self.get_all(DocumentType.CV)
        return docs[0].content_text if docs else None

    async def save_bytes_to_file(self, doc_id: UUID, output_path: str | Path) -> Path:
        """Decode and save a document's binary content to a file."""
        doc = await self.get_by_id(doc_id)
        if not doc or not doc.content_bytes:
            raise ValueError(f"Document {doc_id} not found or has no binary content")
        output_path = Path(output_path)
        output_path.write_bytes(base64.b64decode(doc.content_bytes))
        return output_path

    async def find_similar(self, query: str, top_k: int = 5) -> List[Document]:
        """Find documents most semantically similar to a query string."""
        from sqlalchemy import text

        query_embedding = embed_text(query)
        embedding_str = "[" + ",".join(str(v) for v in query_embedding) + "]"

        async with get_session() as session:
            result = await session.execute(
                text(
                    """
                    SELECT id
                    FROM documents
                    WHERE embedding IS NOT NULL
                    ORDER BY embedding <=> :emb::vector
                    LIMIT :k
                    """
                ),
                {"emb": embedding_str, "k": top_k},
            )
            ids = [row[0] for row in result.fetchall()]

        if not ids:
            return []

        async with get_session() as session:
            result = await session.execute(
                select(Document).where(Document.id.in_(ids))
            )
            return list(result.scalars().all())

    async def upload_directory(self, directory: str | Path) -> List[Document]:
        """
        Upload all documents from a directory.
        Auto-detects document types from filenames.
        """
        directory = Path(directory)
        uploaded = []
        for file_path in directory.iterdir():
            if file_path.suffix.lower() not in (".pdf", ".docx", ".doc", ".txt", ".tex"):
                continue
            doc_type = self._guess_type(file_path.name)
            try:
                doc = await self.upload(file_path, doc_type)
                uploaded.append(doc)
            except Exception as e:
                logger.error(f"[DocStore] Failed to upload {file_path.name}: {e}")
        return uploaded

    def _extract_text(self, path: Path) -> str:
        """Extract text from various file formats."""
        suffix = path.suffix.lower()

        if suffix == ".pdf":
            try:
                from cv.pdf_extractor import extract_text_from_pdf
                return extract_text_from_pdf(path)
            except Exception as e:
                logger.warning(f"[DocStore] PDF extract failed: {e}")
                return ""

        if suffix in (".tex", ".latex"):
            try:
                from cv.latex_extractor import extract_text_from_latex
                return extract_text_from_latex(path)
            except Exception as e:
                logger.warning(f"[DocStore] LaTeX extract failed: {e}")
                return ""

        if suffix in (".docx", ".doc"):
            try:
                import docx
                doc = docx.Document(str(path))
                return "\n".join(p.text for p in doc.paragraphs)
            except Exception as e:
                logger.warning(f"[DocStore] DOCX extract failed: {e}")
                return ""

        if suffix == ".txt":
            return path.read_text(encoding="utf-8", errors="replace")

        return ""

    def _guess_type(self, filename: str) -> DocumentType:
        fn = filename.lower()
        if any(kw in fn for kw in ["cv", "lebenslauf", "resume"]):
            return DocumentType.CV
        if any(kw in fn for kw in ["cover", "anschreiben", "motivation"]):
            return DocumentType.COVER_LETTER_TEMPLATE
        if any(kw in fn for kw in ["zeugnis", "certificate", "cert", "diploma", "abschluss"]):
            return DocumentType.CERTIFICATE
        if any(kw in fn for kw in ["degree", "bachelor", "master", "phd", "doktor"]):
            return DocumentType.DEGREE
        if any(kw in fn for kw in ["reference", "referenz", "empfehlung", "arbeitszeugnis"]):
            return DocumentType.WORK_REFERENCE
        return DocumentType.OTHER
