"""
Cover letter exporter — converts text to professional DOCX and PDF.

Uses python-docx for DOCX generation (fully offline, no external services).
Uses WeasyPrint or ReportLab for PDF generation.
"""
from __future__ import annotations

from pathlib import Path
from typing import Optional

from loguru import logger

from core.models import CVProfileSchema, Job


def export_to_docx(
    content: str,
    output_path: str,
    job: Job,
    cv_profile: CVProfileSchema,
) -> Path:
    """
    Export cover letter text to a professional DOCX file.
    Includes sender details, date, and formatted body.
    """
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt, RGBColor
        from docx.oxml.ns import qn
        import datetime

        doc = Document()

        # Page margins
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.2)
            section.right_margin = Inches(1.2)

        # Sender block (top right)
        sender_para = doc.add_paragraph()
        sender_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        sender_run = sender_para.add_run()
        if cv_profile.full_name:
            sender_run.add_break()
            sender_para.add_run(cv_profile.full_name).bold = True
        if cv_profile.location:
            sender_para.add_run(f"\n{cv_profile.location}")
        if cv_profile.email:
            sender_para.add_run(f"\n{cv_profile.email}")
        if cv_profile.phone:
            sender_para.add_run(f"\n{cv_profile.phone}")

        doc.add_paragraph()

        # Date
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        date_para.add_run(datetime.date.today().strftime("%d. %B %Y"))

        doc.add_paragraph()

        # Company address block
        company_para = doc.add_paragraph()
        company_para.add_run(job.company).bold = True
        if job.location:
            company_para.add_run(f"\n{job.location}")

        doc.add_paragraph()

        # Subject line
        subject_para = doc.add_paragraph()
        subject_run = subject_para.add_run(f"Bewerbung als {job.title}")
        subject_run.bold = True

        doc.add_paragraph()

        # Body — split into paragraphs
        for paragraph in content.split("\n\n"):
            paragraph = paragraph.strip()
            if not paragraph:
                continue
            para = doc.add_paragraph()
            para.add_run(paragraph)
            para.paragraph_format.space_after = Pt(6)

        doc.add_paragraph()

        # Signature
        sig_para = doc.add_paragraph()
        sig_para.add_run("Mit freundlichen Grüßen,\n\n")
        if cv_profile.full_name:
            sig_para.add_run(cv_profile.full_name).bold = True

        # Set font throughout
        for para in doc.paragraphs:
            for run in para.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(11)

        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        logger.info(f"[Exporter] DOCX saved: {output_path}")
        return output_path

    except Exception as e:
        logger.error(f"[Exporter] DOCX export failed: {e}")
        raise


def export_to_pdf(
    content: str,
    output_path: str,
    job: Job,
    cv_profile: CVProfileSchema,
) -> Path:
    """
    Export cover letter to PDF.
    Tries WeasyPrint first (best quality), falls back to ReportLab.
    """
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    # First try: generate from DOCX (highest fidelity)
    docx_path = str(output_path).replace(".pdf", ".docx")
    if Path(docx_path).exists():
        try:
            return _pdf_from_docx(docx_path, output_path)
        except Exception as e:
            logger.debug(f"[Exporter] DOCX→PDF failed: {e}")

    # Second try: WeasyPrint via HTML
    try:
        return _pdf_via_weasyprint(content, output_path, job, cv_profile)
    except Exception as e:
        logger.debug(f"[Exporter] WeasyPrint failed: {e}")

    # Last resort: ReportLab
    return _pdf_via_reportlab(content, output_path, job, cv_profile)


def _pdf_from_docx(docx_path: str, output_path: Path) -> Path:
    """Convert DOCX to PDF using LibreOffice (if available)."""
    import subprocess
    result = subprocess.run(
        [
            "libreoffice", "--headless", "--convert-to", "pdf",
            "--outdir", str(output_path.parent), docx_path,
        ],
        capture_output=True, timeout=30,
    )
    if result.returncode == 0:
        logger.info(f"[Exporter] PDF via LibreOffice: {output_path}")
        return output_path
    raise RuntimeError(f"LibreOffice failed: {result.stderr}")


def _pdf_via_weasyprint(
    content: str, output_path: Path, job: Job, cv_profile: CVProfileSchema
) -> Path:
    import datetime
    from weasyprint import HTML

    html = _build_html(content, job, cv_profile)
    HTML(string=html).write_pdf(str(output_path))
    logger.info(f"[Exporter] PDF via WeasyPrint: {output_path}")
    return output_path


def _pdf_via_reportlab(
    content: str, output_path: Path, job: Job, cv_profile: CVProfileSchema
) -> Path:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        leftMargin=3 * cm,
        rightMargin=2.5 * cm,
        topMargin=2.5 * cm,
        bottomMargin=2.5 * cm,
    )
    styles = getSampleStyleSheet()
    story = []

    for para_text in content.split("\n\n"):
        para_text = para_text.strip()
        if para_text:
            story.append(Paragraph(para_text.replace("\n", "<br/>"), styles["Normal"]))
            story.append(Spacer(1, 0.3 * cm))

    doc.build(story)
    logger.info(f"[Exporter] PDF via ReportLab: {output_path}")
    return output_path


def _build_html(content: str, job: Job, cv_profile: CVProfileSchema) -> str:
    import datetime

    sender_info = []
    if cv_profile.full_name:
        sender_info.append(f"<strong>{cv_profile.full_name}</strong>")
    if cv_profile.location:
        sender_info.append(cv_profile.location)
    if cv_profile.email:
        sender_info.append(cv_profile.email)
    if cv_profile.phone:
        sender_info.append(cv_profile.phone)

    paragraphs = "".join(
        f"<p>{para.strip().replace(chr(10), '<br>')}</p>"
        for para in content.split("\n\n")
        if para.strip()
    )

    return f"""<!DOCTYPE html>
<html lang="de">
<head>
<meta charset="utf-8">
<style>
  body {{ font-family: Calibri, Arial, sans-serif; font-size: 11pt; color: #222; line-height: 1.5; }}
  .sender {{ text-align: right; margin-bottom: 2em; }}
  .date {{ margin-bottom: 1em; }}
  .company {{ margin-bottom: 1em; }}
  .subject {{ font-weight: bold; margin-bottom: 1.5em; }}
  p {{ margin-bottom: 0.8em; text-align: justify; }}
  .signature {{ margin-top: 2em; }}
  @page {{ margin: 2.5cm 2.5cm 2.5cm 3cm; size: A4; }}
</style>
</head>
<body>
  <div class="sender">{"<br>".join(sender_info)}</div>
  <div class="date">{datetime.date.today().strftime("%d. %B %Y")}</div>
  <div class="company"><strong>{job.company}</strong><br>{job.location or ""}</div>
  <div class="subject">Bewerbung als {job.title}</div>
  {paragraphs}
  <div class="signature">
    Mit freundlichen Grüßen,<br><br>
    <strong>{cv_profile.full_name or ""}</strong>
  </div>
</body>
</html>"""
